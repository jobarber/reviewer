import os
import re
import tarfile
from collections import defaultdict

import docx
import matplotlib.pyplot as plt
import pandas as pd
import requests
import seaborn as sns
import torch
from docx.text.run import Run
from transformers import pipeline, AutoModelForSequenceClassification, AutoTokenizer

from peerreviewer import DEVICE
from peerreviewer.models.cohesiveness import CohesiveModel
from peerreviewer.models.emotion import EmotionDetector
from peerreviewer.models.novelty import NoveltyDetector

print('DEVICE ==', DEVICE)


class Document:

    def __init__(self, path_to_docx=None, text=None, cohesiveness_boundary=0.4,
                 novelty_boundary=0.6):
        """
        Constructs the Word document attribute. (Can add other formats later.)

        Args:
            path_to_docx: (str) The path to the .docx file.
            text: (str) A raw text file from which to construct the .docx file.
                If the path_to_docx is None and text is not None, we will attempt to
                construct a .docx document from the raw text file. We assume that
                paragraphs are separated by two newlines ('\n\n'). If both are supplied,
                we will append the text paragraphs to the end of the .docx document.
        """
        if path_to_docx is None and text is None:
            raise TypeError('`path_to_docx` or `text` needs to be provided to `Document`.')
        self.document = docx.Document(path_to_docx) if path_to_docx else docx.Document()
        if text is not None:
            for paragraph in text.split('\n\n'):
                self.document.add_paragraph(paragraph)
        self.cohesiveness_boundary = cohesiveness_boundary
        self.novelty_boundary = novelty_boundary

    def _add_plot(self, x, y, title='Academic Theological Flow', xlabel='Sentence Number',
                  ylabel='Academic Theological "Quality"', temp_figpath='/tmp/academic.png',
                  smooth=False, **sns_kwargs):
        # add image
        # get the rolling average for smoother plot
        if smooth:
            if len(y) > 10:
                y = [sum(y[i:i + 10]) / len(y[i:i + 10]) for i in range(len(y))]
        sns.lineplot(x=x, y=y, **sns_kwargs)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.tight_layout()
        plt.savefig(temp_figpath)
        # add to document
        self.document.paragraphs[0].insert_paragraph_before()
        new_run = self.document.paragraphs[0].add_run()
        new_run.add_picture(temp_figpath)
        plt.close()

    def _adjust_runs(self, paragraph, sentence, softmax, label='',
                     comment_template='Low academic theological score: {softmax}.'):
        num_paragraph_runs = len(paragraph.runs)
        for index in range(len(paragraph.runs)):
            run = paragraph.runs[index]
            # See https://stackoverflow.com/a/52745782
            if sentence in run.text:
                before_text = run.text[:run.text.index(sentence)]
                after_text = run.text[run.text.index(sentence) - 1 + len(sentence):]
                for text in [before_text, sentence, after_text]:
                    new_run = paragraph.add_run(text)
                    if text == sentence:
                        new_run.add_comment(comment_template.format(softmax=softmax, label=label))
            else:
                paragraph.add_run(run.text)

        for run in paragraph.runs[:num_paragraph_runs]:
            paragraph._p.remove(run._r)

    def _analyze_academic_theology(self):
        if not os.path.exists('modelcheckpoints/academic_theology_final'):
            if not os.path.exists('modelcheckpoints'):
                os.mkdir('modelcheckpoints')
                print('downloading academic theology model')

            # download directory
            url = 'https://sagemaker-peerreview.s3.us-west-2.amazonaws.com/academic_theology_final.tar.gz'
            target_path = './modelcheckpoints/academic_theology_final.tar.gz'
            response = requests.get(url, stream=True)
            print('response from AWS S3:', response.status_code)
            if response.status_code == 200:
                with open(target_path, 'wb') as f:
                    f.write(response.raw.read())

            # uncompress directory
            with open(target_path, 'rb') as f:
                file_ = tarfile.open(fileobj=f, mode="r|gz")
                file_.extractall(path='./modelcheckpoints')

            print('finished downloading academic theology model')

        academic_model = AutoModelForSequenceClassification.from_pretrained(
            './modelcheckpoints/academic_theology_final').to(DEVICE)
        tokenizer = AutoTokenizer.from_pretrained('./modelcheckpoints/academic_theology_final')
        probas = []
        for paragraph in self.document.paragraphs:
            # For now parse sentences with regex; make more robust later
            sentences = re.findall(r'[^.?!]{5,}?[.?!]', paragraph.text, flags=re.DOTALL)
            for sentence in sentences:
                with torch.no_grad():
                    output = academic_model(**tokenizer(sentence, return_tensors='pt', padding=True).to(DEVICE))
                    softmax = torch.softmax(output['logits'], dim=-1)
                probas.append(softmax[0][1].item())
                if softmax[0][1] >= 0.4:
                    continue
                self._adjust_runs(paragraph, sentence, softmax[0][1],
                                  comment_template='Low academic theological score: {softmax}.')

        self._add_plot(x=list(range(len(probas))), y=probas, title='Academic Theological Flow (Smoothed)',
                       xlabel='Sentence Number', ylabel='Academic Theological "Quality"',
                       temp_figpath='/tmp/academic.png', smooth=True)

    def _analyze_emotion(self):
        emotion_model = EmotionDetector(pretrained_path='bhadresh-savani/distilbert-base-uncased-emotion')
        emotion_probas = defaultdict(list)
        for paragraph in self.document.paragraphs:
            # For now parse sentences with regex; make more robust later
            sentences = re.findall(r'[^.?!]{5,}?[.?!]', paragraph.text, flags=re.DOTALL)
            for sentence in sentences:
                with torch.no_grad():
                    output = emotion_model.predict_label(sentence)
                    forward_output = emotion_model(sentence)
                    index_probas = torch.softmax(forward_output.logits[0], dim=-1).tolist()
                    for index in range(len(index_probas)):
                        emotion_probas[emotion_model.labels[index]].append(index_probas[index])
                if output['probability'] > 0.6:
                    continue
                self._adjust_runs(paragraph, sentence, output["probability"],
                                  output["label"],
                                  comment_template='Detected emotion: {label} ({softmax:.3f})')

        # Smooth and collate all the emotion probas
        probas = []
        labels = []
        sentence_nums = []
        for label in emotion_probas:
            sub_probas = emotion_probas[label]
            flow_span = len(sub_probas) // 16
            if flow_span > 1:
                emotion_probas[label] = [sum(sub_probas[i:i + 20]) / len(sub_probas[i:i + 20])
                                         for i in range(len(sub_probas))]
            for i, proba in enumerate(sub_probas):
                probas.append(proba)
                labels.append(label)
                sentence_nums.append(i)

        self._add_plot(x=sentence_nums, y=probas, title='Emotional Flow (Smoothed)',
                       xlabel='Sentence Number', ylabel='"Height" of Emotion', temp_figpath='/tmp/emotion.png',
                       hue=labels)

    def _analyze_intradocument_novelty(self):
        novelty_model = NoveltyDetector(pretrained_path='sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')

        probas = []
        for paragraph in self.document.paragraphs:
            contexts = [p.text for p in self.document.paragraphs if p.text != paragraph.text]
            if len(contexts) < 1:
                continue
            with torch.no_grad():
                outputs = novelty_model(contexts, paragraph.text)
            for output in outputs:
                if output >= self.novelty_boundary:
                    probas.append(output.item())
                    break
            else:
                paragraph.add_comment(f'This paragraph is novel with respect to this document:'
                                      f' ({min(1 - output.item(), 1.):.3f}).')
                probas.append(output.item())

        self._add_plot(x=list(range(len(probas))), y=probas, title='Flow of Intradocument Novelty (Smoothed)',
                       xlabel='Paragraph Number', ylabel='Novelty',
                       temp_figpath='/tmp/novelty.png', smooth=True)

    def _analyze_paragraph_cohesion(self):
        cohesive_model = CohesiveModel(pretrained_path='all-mpnet-base-v2')
        probas = []
        for paragraph in self.document.paragraphs:
            # For now parse sentences with regex; make more robust later
            sentences = re.findall(r'[^.?!]{5,}?[.?!]', paragraph.text, flags=re.DOTALL)
            if len(sentences) < 2:
                continue
            with torch.no_grad():
                outputs = cohesive_model.predict(sentences)
            for output in outputs:
                probas.append(output['cohesiveness'])
                if output['cohesiveness'] > self.cohesiveness_boundary:
                    continue
                sentence = sentences[output['sentence_index']]
                self._adjust_runs(paragraph, sentence, output["cohesiveness"],
                                  comment_template='Low cohesion with previous sentence: {softmax:.3f}')

        self._add_plot(x=list(range(len(probas))), y=probas, title='Flow of Cohesiveness (Smoothed)',
                       xlabel='Sentence Number', ylabel='Cohesiveness', temp_figpath='/tmp/cohesiveness.png', smooth=True)

    def _summarize(self):
        pipe = pipeline("summarization",
                        model='t5-base',
                        device=0 if torch.cuda.is_available() else -1,
                        use_fast=True)
        # summary_texts = []
        # for i in range(0, len(self.document.paragraphs), 64):
        prediction = pipe('\n'.join(p.text.replace('\n', ' ') for p in self.document.paragraphs),
                          batch_size=12)
        summary_text = prediction[0]['summary_text']
            # summary_texts.append(summary_text)
        # summary_text = ' '.join(summary_texts)
        # if len(summary_texts) > 3:
        #     prediction = pipe('\n'.join(summary_texts))
        #     summary_text = prediction[0]['summary_text']
        summary_text = 'SUMMARY: ' + summary_text[0].upper() + (summary_text[1:] if summary_text else '') + '\n'
        self.document.paragraphs[0].insert_paragraph_before(summary_text)

    def create_analysis(self):
        """
        Run the analyses specified by methods of class.

        Returns:
            None
        """
        print('analyzing emotion')
        self._analyze_emotion()
        print('analyzing paragraph cohesion')
        self._analyze_paragraph_cohesion()
        print('analyzing intradocument novelty')
        self._analyze_intradocument_novelty()
        print('analyzing academic theology')
        self._analyze_academic_theology()
        print('summarizing')
        self._summarize()

    def save(self, path='test.docx'):
        if '/' not in path:
            self.document.save(os.path.join('peerreviewer/evaldocs', path))
        self.document.save(path)
